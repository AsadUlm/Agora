"""
Text Extraction — converts uploaded files to plain text.

Supported formats:
  .txt   — read as UTF-8 / latin-1 (always available, no extra dep)
  .md    — same as .txt; treated as plain markdown
  .pdf   — pypdf (pure-Python, no native binaries)
  .docx  — python-docx
  .csv   — stdlib csv reader → "header: value | header: value" rows
  .json  — stdlib json with pretty-printed flatten

Extractor registry is a plain dict; add new entries as more formats are needed.

Usage:
    text = extract_text(file_bytes, filename="document.pdf")
"""

from __future__ import annotations

import csv
import io
import json
import logging

logger = logging.getLogger(__name__)

# ── Supported MIME / extension map ────────────────────────────────────────────

# Maps lower-case extension → extractor function name
_SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".csv", ".json"}

# Soft caps to keep extractor output bounded (chunker further enforces sizes).
_CSV_MAX_ROWS = 2000
_JSON_MAX_CHARS = 200_000


def supported_extensions() -> frozenset[str]:
    return frozenset(_SUPPORTED_EXTENSIONS)


# ── Extraction errors ─────────────────────────────────────────────────────────

class ExtractionError(Exception):
    """Raised when text cannot be extracted from the file."""


class UnsupportedFileType(ExtractionError):
    """Raised when the file extension is not supported."""


# ── Individual extractors ─────────────────────────────────────────────────────

def _decode_text(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    raise ExtractionError("Could not decode text file with any supported encoding.")


def _extract_txt(data: bytes) -> str:
    return _decode_text(data)


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader  # noqa: PLC0415 — optional dep
    except ImportError as exc:
        raise ExtractionError("pypdf is not installed. Add pypdf to requirements.txt.") from exc

    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages.append(text)
    result = "\n".join(pages)
    if not result.strip():
        raise ExtractionError("PDF appears to contain no extractable text (scanned image PDF?).")
    return result


def _extract_docx(data: bytes) -> str:
    try:
        import docx  # python-docx  # noqa: PLC0415
    except ImportError as exc:
        raise ExtractionError("python-docx is not installed. Add python-docx to requirements.txt.") from exc

    doc = docx.Document(io.BytesIO(data))
    texts: list[str] = []

    for p in doc.paragraphs:
        if p.text.strip():
            texts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    texts.append(cell.text)

    result = "\n".join(texts)
    if not result.strip():
        raise ExtractionError("DOCX document contains no extractable text.")
    return result


def _extract_csv(data: bytes) -> str:
    """Render CSV as readable rows. Headers are kept as field labels.

    Bounded by ``_CSV_MAX_ROWS`` so a 1M-row CSV does not blow the chunker.
    """
    text = _decode_text(data)
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        raise ExtractionError("CSV file is empty.")
    headers = [h.strip() for h in rows[0]]
    body = rows[1:]
    truncated = False
    if len(body) > _CSV_MAX_ROWS:
        body = body[:_CSV_MAX_ROWS]
        truncated = True

    lines: list[str] = []
    if headers:
        lines.append("Columns: " + ", ".join(headers))
    for row in body:
        if not any(cell.strip() for cell in row):
            continue
        if headers and len(row) == len(headers):
            cells = [f"{h}: {v.strip()}" for h, v in zip(headers, row) if v.strip()]
            lines.append(" | ".join(cells))
        else:
            lines.append(" | ".join(c.strip() for c in row if c.strip()))
    if truncated:
        lines.append(f"… (CSV truncated to first {_CSV_MAX_ROWS} rows)")
    result = "\n".join(lines)
    if not result.strip():
        raise ExtractionError("CSV contains no extractable rows.")
    return result


def _extract_json(data: bytes) -> str:
    """Pretty-print JSON to a bounded text body."""
    text = _decode_text(data)
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError as exc:
        raise ExtractionError(f"JSON is invalid: {exc.msg}") from exc

    pretty = json.dumps(parsed, indent=2, ensure_ascii=False, sort_keys=False)
    if len(pretty) > _JSON_MAX_CHARS:
        pretty = pretty[:_JSON_MAX_CHARS] + "\n…(JSON truncated)…"
    if not pretty.strip():
        raise ExtractionError("JSON document is empty.")
    return pretty


# ── Registry ──────────────────────────────────────────────────────────────────

_EXTRACTORS = {
    ".txt":  _extract_txt,
    ".md":   _extract_txt,
    ".pdf":  _extract_pdf,
    ".docx": _extract_docx,
    ".csv":  _extract_csv,
    ".json": _extract_json,
}


# ── Public API ────────────────────────────────────────────────────────────────

def extension_from_filename(filename: str) -> str:
    """Return the lower-case extension including the dot, e.g. '.pdf'."""
    import os
    _, ext = os.path.splitext(filename)
    return ext.lower()


def extract_text(data: bytes, filename: str) -> str:
    """
    Extract plain text from an uploaded file.

    Args:
        data:     Raw file bytes.
        filename: Original filename (used to determine format).

    Returns:
        Extracted plain-text string.

    Raises:
        UnsupportedFileType: Extension is not in the supported set.
        ExtractionError:     Extraction failed (e.g. corrupted file).
    """
    ext = extension_from_filename(filename)
    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        raise UnsupportedFileType(
            f"File type '{ext}' is not supported. "
            f"Supported: {', '.join(sorted(_SUPPORTED_EXTENSIONS))}"
        )
    logger.debug("Extracting text from '%s' (%d bytes)", filename, len(data))
    return extractor(data)
